# -*- coding: utf-8 -*-
"""
docx_builder.py
Ghep cac Block (da co ban dich) thanh file .docx song ngu Anh-Viet.

Nguyen tac trinh bay:
- Moi doan/tieu de tieng Anh -> ngay ben duoi la ban dich tieng Viet (khac mau/in nghieng
  de de phan biet, khong lam roi mat cau truc).
- Giu nguyen: tieu de (Heading style), in dam/in nghieng trong doan van, danh sach
  (numbered/bulleted), bang bieu (bang docx that, khong phai anh), hinh anh (nhung nguyen,
  KHONG dich chu trong anh), so trang (page number field trong footer).
- Font mac dinh: Times New Roman - ho tro day du tieng Viet co dau tren Word.
"""
import io
from typing import List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .extract import Block

FONT_NAME = "Times New Roman"
VI_TRANSLATION_COLOR = RGBColor(0x1F, 0x4E, 0x79)  # xanh dam, de phan biet voi ban goc (den)
BODY_SIZE = 11
HEADING_SIZES = {1: 18, 2: 15, 3: 13}


def _set_run_font(run, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # dam bao font ap dung cho ca ky tu Unicode (dong bo eastasia/cs de tranh loi font tren mot so may)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)


def _add_page_number_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    _set_run_font(run, size=9)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_runs_from_block(paragraph, block: Block, size=BODY_SIZE):
    """Ghi cac Run goc (giu bold/italic tung phan) vao 1 paragraph docx."""
    for r in block.runs:
        if not r.text:
            continue
        run = paragraph.add_run(r.text)
        _set_run_font(run, size=size, bold=r.bold, italic=r.italic)


def _add_translation_paragraph(doc: Document, text: str, size=BODY_SIZE, italic=True, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    _set_run_font(run, size=size, italic=italic, color=VI_TRANSLATION_COLOR)
    return p


def _build_table(doc: Document, block: Block):
    rows = block.table_rows or []
    vi_rows = block.translated_table_rows or [[""] * len(r) for r in rows]
    if not rows:
        return
    n_rows = len(rows)
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = table.cell(ri, ci)
            en_text = row[ci] if ci < len(row) else ""
            vi_text = vi_rows[ri][ci] if ri < len(vi_rows) and ci < len(vi_rows[ri]) else ""
            # xoa paragraph mac dinh rong
            cell.paragraphs[0].text = ""
            p_en = cell.paragraphs[0]
            is_header_row = (ri == 0)
            run_en = p_en.add_run(en_text)
            _set_run_font(run_en, size=10, bold=is_header_row)
            if vi_text:
                p_vi = cell.add_paragraph()
                run_vi = p_vi.add_run(vi_text)
                _set_run_font(run_vi, size=10, italic=True, bold=is_header_row, color=VI_TRANSLATION_COLOR)
    doc.add_paragraph()  # khoang cach sau bang


def _apply_list_style(paragraph, level: int = 0):
    try:
        paragraph.style = "List Bullet"
    except KeyError:
        pass
    paragraph.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)


def build_bilingual_docx(blocks: List[Block], output_path: str, title: str = "Bilingual Document"):
    doc = Document()

    # dat font mac dinh cho style Normal
    normal_style = doc.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = Pt(BODY_SIZE)
    rpr = normal_style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)

    _add_page_number_footer(doc)

    for block in blocks:
        if block.type == "page_break":
            continue  # khong ep ngat trang cung vi tri PDF goc de tranh trang trong thua

        elif block.type == "heading":
            p_en = doc.add_paragraph()
            p_en.paragraph_format.space_before = Pt(14)
            p_en.paragraph_format.space_after = Pt(2)
            size = HEADING_SIZES.get(block.level, 13)
            for r in block.runs:
                if not r.text.strip() and r.text != " ":
                    continue
                run = p_en.add_run(r.text)
                _set_run_font(run, size=size, bold=True)
            if block.translated_text:
                p_vi = doc.add_paragraph()
                p_vi.paragraph_format.space_after = Pt(10)
                run_vi = p_vi.add_run(block.translated_text)
                _set_run_font(run_vi, size=size - 1, bold=True, italic=True, color=VI_TRANSLATION_COLOR)

        elif block.type == "list_item":
            p_en = doc.add_paragraph()
            _apply_list_style(p_en)
            _add_runs_from_block(p_en, block)
            if block.translated_text:
                p_vi = doc.add_paragraph()
                p_vi.paragraph_format.left_indent = Inches(0.5)
                p_vi.paragraph_format.space_after = Pt(8)
                run_vi = p_vi.add_run(block.translated_text)
                _set_run_font(run_vi, italic=True, color=VI_TRANSLATION_COLOR)

        elif block.type == "paragraph":
            p_en = doc.add_paragraph()
            p_en.paragraph_format.space_after = Pt(2)
            _add_runs_from_block(p_en, block)
            if block.translated_text:
                _add_translation_paragraph(doc, block.translated_text)

        elif block.type == "table":
            _build_table(doc, block)

        elif block.type == "image":
            if block.image_bytes:
                try:
                    width_in = min(max(block.image_width_in, 1.0), 6.3)
                    doc.add_picture(io.BytesIO(block.image_bytes), width=Inches(width_in))
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cap.add_run("[Hình ảnh — giữ nguyên bản gốc, không dịch]")
                    _set_run_font(run, size=8, italic=True)
                except Exception:
                    pass

    doc.core_properties.title = title
    doc.save(output_path)
    return output_path
